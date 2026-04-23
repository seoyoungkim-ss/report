from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.util import Inches

from .excel_reader import SheetData
from .word_extractor import SectionProps, TemplateProfile, clear_document_body


def write_report(
    template_profile: TemplateProfile,
    tagged_content: list[tuple[str, str, str]],
    sheet_map: Optional[dict[str, SheetData]] = None,
    include_data_table: bool = False,
    max_table_rows: int = 30,
) -> bytes:
    doc = template_profile.raw_document
    clear_document_body(doc)

    available_styles = _get_available_styles(doc)

    for tag_type, name, text in tagged_content:
        if tag_type == "TABLE":
            if include_data_table and sheet_map and name in sheet_map:
                _add_data_table(doc, sheet_map[name], max_rows=max_table_rows)
            continue

        if tag_type == "STYLE":
            if text:
                _apply_paragraph(doc, name, text, available_styles)

    _apply_section_props(doc, template_profile.section_props)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _get_available_styles(doc) -> set[str]:
    return {style.name for style in doc.styles}


def _apply_paragraph(doc, style_name: str, text: str, available_styles: set[str]) -> None:
    resolved_style = style_name if style_name in available_styles else "Normal"
    try:
        doc.add_paragraph(text, style=resolved_style)
    except Exception:
        doc.add_paragraph(text)


def _apply_section_props(doc, section_props: SectionProps) -> None:
    try:
        section = doc.sections[0]
        section.left_margin = Inches(section_props.left_margin_inches)
        section.right_margin = Inches(section_props.right_margin_inches)
        section.top_margin = Inches(section_props.top_margin_inches)
        section.bottom_margin = Inches(section_props.bottom_margin_inches)
    except Exception:
        pass


def _add_data_table(doc, sheet: SheetData, max_rows: int = 30) -> None:
    df = sheet.df.head(max_rows)
    if df.empty:
        return

    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = str(col)
        run = hdr_cells[i].paragraphs[0].runs
        if run:
            run[0].bold = True

    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, val in enumerate(row):
            row_cells[col_idx].text = "" if val != val else str(val)
